import os
import requests
import csv
import datetime
import pandas as pd
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from utils import log
from selenium.webdriver.common.by import By
from config import OUTPUT_DIR


def export_to_csv(data, start_date, headers=None, label="BOE Account", output_dir=None):
    filename = build_filename(start_date, label, ext="csv", output_dir=output_dir)
    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(data)
    log(f"Exported {len(data)} rows to {filename}")
    return filename


def export_to_excel(data, start_date, headers=None, label="BOE Account", output_dir=None):
    filename = build_filename(start_date, label, ext="xlsx", output_dir=output_dir)
    df = pd.DataFrame(data, columns=headers)
    df.to_excel(filename, index=False)
    log(f"Exported {len(df)} rows to {filename}")
    return filename




# --------------------------
# PDF Export
# --------------------------

# Custom PDF class with footer timestamp
class CustomPDF(FPDF):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.generated_on = datetime.datetime.now().strftime("%d-%m-%Y %H:%M")

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Generated on {self.generated_on}", align="C")


def ensure_font():
    font_dir = os.path.join(os.path.dirname(__file__), "fonts")
    os.makedirs(font_dir, exist_ok=True)
    font_path = os.path.join(font_dir, "NotoSans-Regular.ttf")

    if not os.path.exists(font_path):
        try:
            url = "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf"
            r = requests.get(url, timeout=10)
            r.raise_for_status()
            with open(font_path, "wb") as f:
                f.write(r.content)
        except Exception:
            # Fallback to system font (Windows example: Arial)
            system_font = r"C:\Windows\Fonts\arial.ttf"
            if os.path.exists(system_font):
                return system_font
            else:
                raise RuntimeError("No font available")
    return font_path

def export_to_pdf(data, start_date, headers=None, label="BOE Account", output_dir=None):
    filename = build_filename(start_date, label, ext="pdf", output_dir=output_dir)
    pdf = CustomPDF(orientation="L")
    pdf.add_page()

    font_path = ensure_font()
    pdf.add_font("CustomFont", "", font_path)
    pdf.add_font("CustomFont", "B", font_path)

    page_width = pdf.w - 2 * pdf.l_margin
    col_count = len(headers) if headers else len(data[0])
    col_width = page_width / col_count
    row_height = 8

    if headers:
        pdf.set_font("CustomFont", "B", size=11)
        for h in headers:
            pdf.cell(col_width, row_height, h, border=1, align="C")
        pdf.ln(row_height)

    pdf.set_font("CustomFont", "", size=10)
    for r_index, row in enumerate(data):
        fill = (r_index % 2 == 0)
        pdf.set_fill_color(230, 230, 230) if fill else pdf.set_fill_color(255, 255, 255)
        for cell_text in row:
            align = "R" if cell_text.replace(".", "").isdigit() else "L"
            pdf.cell(col_width, row_height, cell_text, border=1, align=align, fill=True)
        pdf.ln(row_height)

    pdf.set_font("CustomFont", "B", size=11)
    pdf.set_fill_color(200, 200, 200)
    pdf.cell(page_width, row_height, f"Total rows: {len(data)}",
         border=1, align="R", fill=True,
         new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(filename)
    log(f"Exported {len(data)} rows to {filename}")
    return filename





# --------------------------
# Word / DOCX Export
# --------------------------

def set_cell_shading(cell, fill="E6E6E6"):
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.append(tcPr)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def export_to_docx(data, start_date, headers=None, label="BOE Account", output_dir=None):
    filename = build_filename(start_date, label, ext="docx", output_dir=output_dir)
    doc = Document()

    # Force landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading(f"{label} Report", level=1)

    col_count = len(headers) if headers else len(data[0])
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers or []):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows with alternating shading
    for r_index, row in enumerate(data):
        row_cells = table.add_row().cells
        for c_index, cell_text in enumerate(row):
            run = row_cells[c_index].paragraphs[0].add_run(cell_text)
            run.font.size = Pt(10)

        if r_index % 2 == 0:
            for cell in row_cells:
                set_cell_shading(cell, fill="E6E6E6")  # light gray

    # Summary row with darker shading + bold emphasis
    summary_row = table.add_row().cells
    run = summary_row[-1].paragraphs[0].add_run(f"Total rows: {len(data)}")
    run.bold = True
    run.font.size = Pt(11)
    for cell in summary_row:
        set_cell_shading(cell, fill="C8C8C8")  # darker gray
        # Make summary row text bold for emphasis
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True

    # Footer with timestamp
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    timestamp = datetime.datetime.now().strftime("%d-%m-%Y %H:%M")
    p.text = f"Generated on {timestamp}"

    doc.save(filename)
    log(f"Exported {len(data)} rows to {filename}")
    return filename



# --------------------------
# Master Export Function
# --------------------------
def export_all_formats(data, headers=None, report_rows=None, start_date=None, label="BOE Account", output_dir=None):
    csv_file = export_to_csv(data, start_date, headers=headers, label=label, output_dir=output_dir)
    xlsx_file = export_to_excel(data, start_date, headers=headers, label=label, output_dir=output_dir)
    pdf_file = export_to_pdf(report_rows or data, start_date, headers=headers, label=label, output_dir=output_dir)
    docx_file = export_to_docx(report_rows or data, start_date, headers=headers, label=label, output_dir=output_dir)
    return {
        "csv": csv_file,
        "xlsx": xlsx_file,
        "pdf": pdf_file,
        "docx": docx_file,
    }


def build_filename(start_date, label="BOE Account", ext="csv", output_dir=None):
    dt = datetime.datetime.strptime(start_date, "%d/%m/%Y")
    month = dt.strftime("%B")
    year = dt.strftime("%Y")
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{month} {year} {label}_{timestamp}.{ext}"

    # Use explicit output_dir if passed, else fallback to config.OUTPUT_DIR
    target_dir = output_dir or OUTPUT_DIR
    if target_dir:
        os.makedirs(target_dir, exist_ok=True)
        return os.path.join(target_dir, filename)
    return filename


def scrape_headers(driver):
    headers = []
    try:
        header_elements = driver.find_elements(By.CSS_SELECTOR, "table.g-table thead th")
        headers = [h.text.strip() for h in header_elements]
    except Exception as e:
        log(f"Failed to scrape headers: {e}")
    return headers
